import os
import sys

# Auto-install python-docx if missing
try:
    import docx
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

def main():
    template_path = r"d:\Gx-RollNumber_AI_Usage_Report.docx"
    output_path = r"d:\G3-HE204327_AI_Usage_Report.docx"
    
    print("Restoring reports back to their original template states...")
    
    for path in [template_path, output_path]:
        if not os.path.exists(path):
            continue
        try:
            doc = docx.Document(path)
            
            # Revert title
            for p in doc.paragraphs:
                if "PROJECT ASSIGNMENT: EDUNEXUS LEARNING PLATFORM" in p.text:
                    p.text = "PROJECT ASSIGNMENT/LAB"
                if "AI for SDLC - Usage Report" in p.text:
                    p.text = "AI for sdlc - Usage report"
                    
            # Revert General Info
            general_reversions = {
                "Nguyễn Quang Huy - HE204327 (Group G3)": "[Tên, MSSV]",
                "Claude 3.5 Sonnet / Antigravity AI IDE Coding Assistant (chat), ChatGPT-4o": "[Tên + phiên bản/model, ví dụ: Claude Sonnet 4.6 (chat), GitHub Copilot (VS Code), ChatGPT-4o; nếu dùng các AI các nhau cho các output khác nhau thì cũng đề cập cụ thể dùng AI nào cho output gì]",
                "Phase 1 - Requirements Analysis (Updating Screen Flow & Screen Inventory)": "[Toàn bộ project / Giai đoạn cụ thể]",
                "ai-log.md (located in repository root)": "[Export conversation / file ai-log.md / screenshot - đính kèm hoặc link]",
                "Used Claude 3.5 Sonnet to design the system mappings and write python automation scripts to parse HTML templates and inject them into FDS docx files.": "[Mô tả ngắn: có dùng AI Project/workspace để giữ context xuyên suốt? Có 1 hay nhiều phiên AI riêng cho mỗi giai đoạn?]"
            }
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for old_txt, new_txt in general_reversions.items():
                                if old_txt in p.text:
                                    p.text = p.text.replace(old_txt, new_txt)
                                    
                        # Revert Contribution Map rows
                        if "Screen Flow Diagram (EduNexus_ScreenFlow_Updated.drawio)" in cell.text:
                            row.cells[0].paragraphs[0].text = "[ví dụ: UCS / UC-05 Submit Application]"
                            row.cells[1].paragraphs[0].text = "[B]"
                            row.cells[2].paragraphs[0].text = '[Sửa lại Alternative Flow AF-03 vì AI bỏ sót case "đã nộp đơn rồi"; thêm BR-02 giới hạn 5MB theo đề bài]'
                        elif "Screen Inventory Specification (Screen_Inventory.md)" in cell.text:
                            row.cells[0].paragraphs[0].text = "[ví dụ: TDS / 3.Data Model]"
                            row.cells[1].paragraphs[0].text = "[A]"
                            row.cells[2].paragraphs[0].text = "-"
                        elif "FDS Document Generator Script (generate_fds.py)" in cell.text:
                            row.cells[0].paragraphs[0].text = "[ví dụ: Service ApplicationService]"
                            row.cells[1].paragraphs[0].text = "[C]"
                            row.cells[2].paragraphs[0].text = "[SV viết khung method + logic chính, AI hỗ trợ viết Javadoc và xử lý exception]"
                        elif cell.text == "-":
                            # Revert additional rows
                            row.cells[0].paragraphs[0].text = "[Thêm dòng cho mỗi artifact chính của giai đoạn]"
                            row.cells[1].paragraphs[0].text = ""
                            row.cells[2].paragraphs[0].text = ""
                            
                        # Revert AI Errors row
                        if "Vẽ sơ đồ Screen Flow" in cell.text:
                            row.cells[0].paragraphs[0].text = "[ví dụ: sinh UC Index]"
                            row.cells[1].paragraphs[0].text = '[AI thêm Actor "Recruiter" không có trong đề bài]'
                            row.cells[2].paragraphs[0].text = "[Đề bài chỉ có 4 actor: Candidate/HR/Interviewer/Admin]"
                            row.cells[3].paragraphs[0].text = "[Yêu cầu AI loại bỏ, không thêm UC liên quan]"
                            
            # Revert paragraphs
            for p in doc.paragraphs:
                if '"đây là file màn hình thầy bảo làm và template mới hãy vào D:\\stitch_edunexus_learning_platform (1)\\stitch_edunexus_learning_platform để sửa lại phần 1. Screen Inventory đi (sửa lại bằng tiếng anh đi)"' in p.text:
                    p.text = "[Prompt đã dùng: nguyên văn]"
                if '"I have designed the 26-screen mapping for EduNexus based on the Excel sheet and created generate_fds.py to parse the HTML and insert them into the Word document... (truncated)... Screen_Inventory.md generated successfully in English."' in p.text:
                    p.text = '[Output AI trả về: nguyên văn hoặc rút gọn, ghi rõ "...(cắt bớt)..." nếu dài]'
                if "Nhận xét: AI hiểu đúng yêu cầu, phân loại chính xác các màn hình thành 6 nhóm tính năng chính. Tuy nhiên, do môi trường Sandbox bị chặn chạy lệnh terminal trên ổ D, AI đã chuyển hướng cung cấp script tự động hóa python-docx để sinh viên tự chạy cục bộ, đây là một giải pháp rất linh hoạt và thông minh." in p.text:
                    p.text = "Nhận xét: [AI có hiểu đúng yêu cầu không? Có cần hỏi lại/làm rõ thêm không? Output này sau đó được dùng như thế nào - giữ nguyên, sửa, hay bỏ?]"
                    
                if "- Screen Flow Diagram: Thay đổi liên kết màn hình Profile và Password Change nối thẳng vào Login. Xóa bỏ các nút chức năng nhỏ màu xám xung quanh các màn hình.\n- Screen Inventory: Dịch toàn bộ sang tiếng Anh, gộp màn hình Flashcard Library và bổ sung đầy đủ mô tả thuộc tính input (Required, Placeholder)." in p.text:
                    p.text = "Đã thay đổi gì? [Mô tả cụ thể - thêm/xóa/sửa phần nào]"
                if "Vì sao thay đổi? Chọn 1 hoặc nhiều lý do và giải thích:" in p.text:
                    if "AI sai/thiếu so với đề bài" in p.text:
                        p.text = "Vì sao thay đổi? Chọn 1 hoặc nhiều lý do và giải thích:\n☐ AI sai/thiếu so với đề bài hoặc tài liệu giai đoạn trước (nêu rõ tài liệu/dòng nào)\n☐ AI đúng về kỹ thuật nhưng không phù hợp domain (giải thích vì sao lại không)\n☐ AI viết quá phức tạp/quá đơn giản so với scope môn học\n☐ AI vi phạm convention/template đã định\n☐ Khác: [ghi rõ]"
                        
                if "Q1: Giải thích sự khác biệt giữa Course Structure (SCR-05) và Student Dashboard (SCR-02) trong hệ thống EduNexus." in p.text:
                    p.text = '[Câu hỏi 1 - ví dụ: "Giải thích Business Rule quan trọng nhất của UC-05 và vì sao nó cần thiết với đề bài này"]'
                if "Q2: Tại sao Assignment Submit (SCR-12) and Assignment Result (SCR-13) lại được tách biệt thành hai màn hình riêng?" in p.text:
                    p.text = '[Câu hỏi 2 - ví dụ: "Vì sao Entity X có quan hệ N-N với Entity Y? Nếu bỏ bảng trung gian thì hệ thống gặp vấn đề gì?"]'
                    
                if "Việc nào AI làm tốt hơn nhóm tự làm? Vì sao?" in p.text:
                    if "DOM" in p.text or "Word" in p.text:
                        p.text = "Việc nào AI làm tốt hơn nhóm tự làm? Vì sao?\n[Việc nào AI làm tốt hơn nhóm tự làm? Vì sao?]"
                if "Việc nào nhóm làm tốt hơn AI / AI không thể làm tốt?" in p.text:
                    if "Sơ đồ" in p.text or "SME" in p.text:
                        p.text = "Việc nào nhóm làm tốt hơn AI / AI không thể làm tốt? Vì sao (kiến thức domain, ràng buộc đề bài, kinh nghiệm thực tế...)?\n[Việc nào nhóm làm tốt hơn AI / AI không thể làm tốt? Vì sao (kiến thức domain, ràng buộc đề bài, kinh nghiệm thực tế...)?]"
                if "Rủi ro lớn nhất nếu KHÔNG review kỹ output của AI" in p.text:
                    if "FDS" in p.text or "Use Case" in p.text:
                        p.text = "Rủi ro lớn nhất nếu KHÔNG review kỹ output của AI trong project này là gì? Cho 1 ví dụ cụ thể đã/suýt xảy ra.\n[Rủi ro lớn nhất nếu KHÔNG review kỹ output của AI trong project này là gì? Cho 1 ví dụ cụ thể đã/suýt xảy ra.]"
                if "Nếu làm lại từ đầu, nhóm sẽ thay đổi gì trong cách dùng AI" in p.text:
                    if "Glossary" in p.text or "Design System" in p.text:
                        p.text = "Nếu làm lại từ đầu, nhóm sẽ thay đổi gì trong cách dùng AI (ví dụ: prompt khác, thứ tự khác, công cụ khác)?\n[Nếu làm lại từ đầu, nhóm sẽ thay đổi gì trong cách dùng AI (ví dụ: prompt khác, thứ tự khác, công cụ khác)?]"
                        
                if "- File ai-log.md (located in the repository root directory) contains the timeline log of all prompts, AI outputs, and manual changes." in p.text:
                    p.text = "File export conversation đầy đủ (Claude/ChatGPT share link hoặc file .json/.md export), HOẶC"
                if "- The scripts generate_fds.py, complete_report.py, and read_report_template.py are also included in the repository." in p.text:
                    p.text = "File ai-log.md ghi các prompt/lệnh chính theo timeline, kèm timestamp/commit hash tương ứng."
                    
            doc.save(path)
            print(f"Successfully restored: {path}")
        except Exception as e:
            print(f"Error restoring {path}: {e}")

if __name__ == "__main__":
    main()
