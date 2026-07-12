import os
import re
import sys
import shutil

# Check docx dependency and install if missing
try:
    import docx
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

# Try beautifulsoup4 for better HTML parsing, fallback to regex if not available
try:
    from bs4 import BeautifulSoup
    USE_BS4 = True
except ImportError:
    print("BeautifulSoup4 not installed. Falling back to regex parsing.")
    USE_BS4 = False

# Mapping of directories to Screen IDs, Names, Roles, and Purpose (English Version)
SCREENS_CONFIG = [
    {
        "id": "SCR-01",
        "name": "User Login",
        "group": "Common",
        "role": "Student / SME",
        "dirs": ["edunexus_login"],
        "purpose": "Google Login and credentials login for Students and SMEs."
    },
    {
        "id": "SCR-02",
        "name": "Student Dashboard",
        "group": "Common",
        "role": "Student",
        "dirs": ["student_dashboard"],
        "purpose": "Student main homepage displaying registered courses, classes, and content packages."
    },
    {
        "id": "SCR-03",
        "name": "Personal Progress",
        "group": "Common",
        "role": "Student",
        "dirs": ["personal_progress"],
        "purpose": "Displays detailed study metrics, completed lessons, assignments, and test history."
    },
    {
        "id": "SCR-04",
        "name": "Course List",
        "group": "Common",
        "role": "SME",
        "dirs": ["course_list_sme"],
        "purpose": "SME dashboard containing courses assigned for design and content management."
    },
    {
        "id": "SCR-05",
        "name": "Course Structure",
        "group": "Common",
        "role": "SME",
        "dirs": ["course_structure_editor"],
        "purpose": "Course syllabus tree builder where SME configures chapters, lessons, quizzes, assignments, and flashcards."
    },
    {
        "id": "SCR-06",
        "name": "Lesson Editor",
        "group": "Lesson",
        "role": "SME",
        "dirs": ["lesson_editor"],
        "purpose": "Workspace for SME to compose lesson content, add lecture video links, and upload resource documents."
    },
    {
        "id": "SCR-07",
        "name": "AI Lesson Staging",
        "group": "Lesson",
        "role": "SME",
        "dirs": ["ai_lesson_staging"],
        "purpose": "Staging console where SME reviews, edits, and approves AI-generated lesson content."
    },
    {
        "id": "SCR-08",
        "name": "Lesson Text Extract",
        "group": "Lesson",
        "role": "SME",
        "dirs": ["lesson_text_extract"],
        "purpose": "YouTube video scraper that extracts transcripts and leverages GenAI to formulate lesson summaries."
    },
    {
        "id": "SCR-09",
        "name": "Lesson View",
        "group": "Lesson",
        "role": "Student",
        "dirs": ["lesson_view_student"],
        "purpose": "Student viewport for lessons, offering video playback, reading content, and resource downloads."
    },
    {
        "id": "SCR-10",
        "name": "Assignment List",
        "group": "Assignment",
        "role": "SME",
        "dirs": ["assignment_list_sme"],
        "purpose": "Listing dashboard for SME to monitor, grade, and organize essay assignments."
    },
    {
        "id": "SCR-11",
        "name": "Assignment Detail",
        "group": "Assignment",
        "role": "SME",
        "dirs": ["assignment_detail_sme"],
        "purpose": "Setting details and defining evaluation rubrics for essay assignments."
    },
    {
        "id": "SCR-12",
        "name": "Assignment Submit",
        "group": "Assignment",
        "role": "Student",
        "dirs": ["assignment_submit_student"],
        "purpose": "Submitting text essays or uploading file attachments for assignments."
    },
    {
        "id": "SCR-13",
        "name": "Assignment Result",
        "group": "Assignment",
        "role": "Student",
        "dirs": ["assignment_result_student"],
        "purpose": "Displaying grading details, scored criteria, and AI feedback."
    },
    {
        "id": "SCR-14",
        "name": "Flashcard Editor",
        "group": "Flashcard",
        "role": "SME",
        "dirs": ["flashcard_editor_sme"],
        "purpose": "SMEs build vocabulary or concept cards by typing Side A and Side B text."
    },
    {
        "id": "SCR-15",
        "name": "AI Flashcard Staging",
        "group": "Flashcard",
        "role": "SME",
        "dirs": ["ai_flashcard_staging"],
        "purpose": "Editing and filtering flashcard decks generated automatically by AI from lessons."
    },
    {
        "id": "SCR-16",
        "name": "Flashcard Library",
        "group": "Flashcard",
        "role": "Student",
        "dirs": ["flashcard_library_student_1", "flashcard_library_student_2"],
        "purpose": "Collection of student's personal decks showing completion and mastery status."
    },
    {
        "id": "SCR-17",
        "name": "Flashcard Practice",
        "group": "Flashcard",
        "role": "Student",
        "dirs": ["flashcard_practice_student"],
        "purpose": "Interactive slide deck showing cards with 3D rotation and spaced repetition buttons."
    },
    {
        "id": "SCR-18",
        "name": "Question Bank",
        "group": "Question",
        "role": "SME",
        "dirs": ["question_bank_sme"],
        "purpose": "Centralized index where SME coordinates course test questions."
    },
    {
        "id": "SCR-19",
        "name": "Question Detail",
        "group": "Question",
        "role": "SME",
        "dirs": ["question_detail_sme"],
        "purpose": "Formatting questions and options while specifying the correct solution and explanations."
    },
    {
        "id": "SCR-20",
        "name": "AI Question Staging",
        "group": "Question",
        "role": "SME",
        "dirs": ["ai_question_staging_sme"],
        "purpose": "Editing and filtering multiple choice test questions generated by AI."
    },
    {
        "id": "SCR-21",
        "name": "Question Import",
        "group": "Question",
        "role": "SME",
        "dirs": ["question_import_sme"],
        "purpose": "Bulk importing test questions using template Excel formatting."
    },
    {
        "id": "SCR-22",
        "name": "Quiz History",
        "group": "Quiz",
        "role": "Student",
        "dirs": ["quiz_history_student"],
        "purpose": "Student summary page highlighting test records and metrics."
    },
    {
        "id": "SCR-23",
        "name": "New Quiz",
        "group": "Quiz",
        "role": "Student",
        "dirs": ["new_quiz_student"],
        "purpose": "Configuring and starting quiz practice sessions."
    },
    {
        "id": "SCR-24",
        "name": "Quiz Taking",
        "group": "Quiz",
        "role": "Student",
        "dirs": ["quiz_taking_student"],
        "purpose": "Timer-based testing workspace with question navigation list."
    },
    {
        "id": "SCR-25",
        "name": "Quiz Results",
        "group": "Quiz",
        "role": "Student",
        "dirs": ["quiz_results_student"],
        "purpose": "Quiz summary statistics detailing scores, time elapsed, and correct counts."
    },
    {
        "id": "SCR-26",
        "name": "Quiz Review",
        "group": "Quiz",
        "role": "Student",
        "dirs": ["quiz_review_student"],
        "purpose": "Question-by-question response review displaying correct answers and explanations."
    }
]

def clean_text(text):
    if not text:
        return ""
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

def parse_html_bs4(html_content):
    soup = BeautifulSoup(html_content, 'html.parser')
    components = []
    
    # Identify title or header
    title_el = soup.find('title')
    title = title_el.text.strip() if title_el else "EduNexus Screen"
    
    # Find all major interactive elements
    # 1. Inputs
    for input_el in soup.find_all(['input', 'select', 'textarea']):
        el_id = input_el.get('id', input_el.get('name', ''))
        el_type = input_el.get('type', input_el.name)
        
        # Try to find associated label
        label_text = ""
        if el_id:
            label_el = soup.find('label', attrs={'for': el_id})
            if label_el:
                label_text = label_el.text.strip()
        
        if not label_text:
            # Check parent label or nested text
            parent = input_el.parent
            if parent and parent.name == 'label':
                label_text = parent.text.strip()
                
        # Fallback to placeholder or name
        placeholder = input_el.get('placeholder', '')
        desc = f"Input field: {label_text}" if label_text else f"Input field"
        if placeholder:
            desc += f", placeholder: \"{placeholder}\""
        if input_el.get('required') is not None:
            desc += " (Required)"
            
        name = label_text if label_text else (placeholder if placeholder else el_id)
        if not name:
            name = f"{el_type.capitalize()} field"
            
        components.append({
            "name": name,
            "type": f"Input {el_type}",
            "desc": desc
        })
        
    # 2. Buttons
    for btn_el in soup.find_all(['button', 'input'], attrs={'type': ['submit', 'button']}):
        if btn_el.name == 'input' and btn_el.get('type') not in ['submit', 'button']:
            continue
        btn_text = btn_el.text.strip() if btn_el.name == 'button' else btn_el.get('value', '')
        if not btn_text:
            btn_text = btn_el.get('aria-label', btn_el.get('id', 'Button'))
            
        btn_text = clean_text(btn_text)
        components.append({
            "name": f"\"{btn_text}\" button",
            "type": "Button",
            "desc": f"Click to execute action: {btn_text}"
        })
        
    # 3. Regular buttons without type
    for btn_el in soup.find_all('button'):
        if btn_el.get('type') in ['submit', 'button']:
            continue # already processed
        btn_text = clean_text(btn_el.text)
        if not btn_text:
            btn_text = btn_el.get('aria-label', btn_el.get('id', 'Button'))
        components.append({
            "name": f"\"{btn_text}\" button",
            "type": "Button",
            "desc": f"Click to trigger action/open popup: {btn_text}"
        })

    # 4. Action Links
    for a_el in soup.find_all('a'):
        a_text = clean_text(a_el.text)
        href = a_el.get('href', '#')
        if not a_text or a_text.startswith('http') or len(a_text) > 50:
            continue
        
        components.append({
            "name": f"\"{a_text}\" link",
            "type": "Text link",
            "desc": f"Redirect link: {a_text} (Destination: {href})"
        })
        
    # 5. Tables or Lists
    for tbl in soup.find_all('table'):
        headers = [th.text.strip() for th in tbl.find_all('th') if th.text.strip()]
        tbl_name = "Data table"
        if headers:
            tbl_name = f"Data table ({', '.join(headers[:3])}...)"
        components.append({
            "name": tbl_name,
            "type": "Table",
            "desc": "Displays structured data list."
        })

    return title, components

def parse_html_regex(html_content):
    components = []
    
    # Page Title
    title_match = re.search(r'<title>(.*?)</title>', html_content, re.IGNORECASE | re.DOTALL)
    title = title_match.group(1).strip() if title_match else "EduNexus Screen"
    
    # Inputs
    input_matches = re.finditer(r'<input\s+([^>]*?)>', html_content, re.IGNORECASE)
    for m in input_matches:
        attrs = m.group(1)
        el_type = re.search(r'type=["\'](.*?)["\']', attrs, re.IGNORECASE)
        el_id = re.search(r'id=["\'](.*?)["\']', attrs, re.IGNORECASE)
        placeholder = re.search(r'placeholder=["\'](.*?)["\']', attrs, re.IGNORECASE)
        required = 'required' in attrs
        
        t = el_type.group(1) if el_type else "text"
        i = el_id.group(1) if el_id else ""
        p = placeholder.group(1) if placeholder else ""
        
        if t in ['checkbox', 'radio', 'text', 'email', 'password', 'number', 'file']:
            name = p if p else (i if i else f"Input {t}")
            desc = f"Input field {t}"
            if p:
                desc += f" (placeholder: \"{p}\")"
            if required:
                desc += " - Required"
            components.append({
                "name": name,
                "type": f"Input {t}",
                "desc": desc
            })
            
    # Buttons
    btn_matches = re.finditer(r'<button\s*[^>]*?>(.*?)</button>', html_content, re.IGNORECASE | re.DOTALL)
    for m in btn_matches:
        text = clean_text(re.sub(r'<[^>]*?>', '', m.group(1)))
        if not text:
            text = "Button"
        if len(text) < 40:
            components.append({
                "name": f"\"{text}\" button",
                "type": "Button",
                "desc": f"Execute action: {text}"
            })
            
    # Links
    link_matches = re.finditer(r'<a\s+[^>]*?href=["\'](.*?)["\'][^>]*?>(.*?)</a>', html_content, re.IGNORECASE | re.DOTALL)
    for m in link_matches:
        href = m.group(1)
        text = clean_text(re.sub(r'<[^>]*?>', '', m.group(2)))
        if text and not text.startswith('http') and len(text) < 40:
            components.append({
                "name": f"\"{text}\" link",
                "type": "Text link",
                "desc": f"Redirect to {href}"
            })
            
    return title, components

def analyze_navigation_and_conditions(screen_id, screen_name, components):
    # Determine navigation and conditions logically in English
    nav_flows = []
    display_conds = []
    
    if screen_id == "SCR-01":
        nav_flows = [
            "Successful login as Student → Redirect to Student Dashboard (SCR-02).",
            "Successful login as SME → Redirect to Course List (SCR-04).",
            "Click \"Register now\" → Redirect to Registration page (if available).",
            "Click \"Forgot password?\" → Display reset password confirmation."
        ]
        display_conds = [
            "If the user is already authenticated, automatically redirect to the respective role dashboard.",
            "Password text is masked by default with show/hide toggle.",
            "Show error banner when credentials are invalid or if the account is locked."
        ]
    elif screen_id == "SCR-02":
        nav_flows = [
            "Click on any course card → Redirect to Lesson View (SCR-09).",
            "Click on learning progress → Redirect to Personal Progress (SCR-03).",
            "Click on Flashcards option → Redirect to Flashcard Library (SCR-16).",
            "Click \"Quick Practice Test\" → Redirect to New Quiz (SCR-23)."
        ]
        display_conds = [
            "Requires role: Student.",
            "Enrolled courses list is automatically populated based on registered or paid courses of the active student."
        ]
    elif screen_id == "SCR-03":
        nav_flows = [
            "Click \"Back to Dashboard\" → Redirect to Student Dashboard (SCR-02).",
            "Click \"Continue Learning\" → Redirect to Lesson View (SCR-09) of the first uncompleted lesson."
        ]
        display_conds = [
            "Display progress bars as percentages (%).",
            "Data is loaded dynamically based on real-time activity tracking of the authenticated student."
        ]
    elif screen_id == "SCR-04":
        nav_flows = [
            "Click on a course name or \"Edit Structure\" → Redirect to Course Structure (SCR-05).",
            "Click logout → Redirect to Login (SCR-01)."
        ]
        display_conds = [
            "Requires role: SME.",
            "Only displays courses where the active SME is assigned authoring/management rights."
        ]
    elif screen_id == "SCR-05":
        nav_flows = [
            "Click on any lesson name → Open Lesson Editor (SCR-06).",
            "Click \"AI Gen\" option → Redirect to AI Lesson Staging (SCR-07) or AI Flashcard Staging (SCR-15).",
            "Click on any assignment name → Open Assignment Detail (SCR-11).",
            "Click on Question Bank → Open Question Bank (SCR-18)."
        ]
        display_conds = [
            "Requires role: SME.",
            "Syllabus is displayed as a hierarchical tree structure (Modules > Chapters > Activities).",
            "Save button activates only when structural changes are detected."
        ]
    elif screen_id == "SCR-06":
        nav_flows = [
            "Click \"Back to Course\" → Redirect to Course Structure (SCR-05).",
            "Click \"Extract from YouTube\" → Redirect to Lesson Text Extract (SCR-08)."
        ]
        display_conds = [
            "Requires role: SME.",
            "Built-in editor supports live split-pane side-by-side Markdown rendering.",
            "Autosaves content drafts every 30 seconds."
        ]
    elif screen_id == "SCR-07":
        nav_flows = [
            "Click \"Approve & Save\" → Save and redirect to Course Structure (SCR-05).",
            "Click \"Cancel\" → Return to Course Structure (SCR-05) without changes."
        ]
        display_conds = [
            "Requires role: SME.",
            "Screen contains split-column views comparing original AI draft and editable workspace side-by-side.",
            "Approve action is disabled if the editor content is empty."
        ]
    elif screen_id == "SCR-08":
        nav_flows = [
            "Successful extraction → Automatically redirects and populates the text in AI Lesson Staging (SCR-07)."
        ]
        display_conds = [
            "Requires role: SME.",
            "Input string must be a valid YouTube URL. Shows error alert if scraping fails or video lacks subtitles."
        ]
    elif screen_id == "SCR-09":
        nav_flows = [
            "Click \"Next Lesson\" → Redirect to next lesson in sequence.",
            "Click attached activity → Redirect to Assignment Submit (SCR-12) or Quiz Taking (SCR-24)."
        ]
        display_conds = [
            "Requires role: Student.",
            "Marks completion status automatically when the video finishes or reader scrolls to bottom of content."
        ]
    elif screen_id == "SCR-10":
        nav_flows = [
            "Click on assignment title or \"Edit\" → Redirect to Assignment Detail (SCR-11).",
            "Click \"Add New\" → Redirect to Assignment Detail (SCR-11) with empty form fields."
        ]
        display_conds = [
            "Requires role: SME.",
            "Lists all assignments with counts of submitted and graded essays."
        ]
    elif screen_id == "SCR-11":
        nav_flows = [
            "Click \"Save\" → Save configurations and redirect to Assignment List (SCR-10)."
        ]
        display_conds = [
            "Requires role: SME.",
            "Grid supports adding multi-level evaluation criteria. Combined rubric weight parameters must equal 100%."
        ]
    elif screen_id == "SCR-12":
        nav_flows = [
            "Click \"Submit\" → Submits work and redirects to Assignment Result (SCR-13) for AI grading."
        ]
        display_conds = [
            "Requires role: Student.",
            "Submit options block automatically when the deadline expires.",
            "File attachment upload size is limited to 10MB per file."
        ]
    elif screen_id == "SCR-13":
        nav_flows = [
            "Click \"Back to Course\" → Return to Lesson View (SCR-09) or Course Structure."
        ]
        display_conds = [
            "Requires role: Student.",
            "Render final score and rubric table containing detailed score mappings clearly.",
            "Render AI evaluation text in a readable rich-text format."
        ]
    elif screen_id.startswith("SCR-14"):
        nav_flows = [
            "Click \"Save\" → Redirect back to Course Structure (SCR-05)."
        ]
        display_conds = [
            "Requires role: SME.",
            "Decks support unlimited card capacities."
        ]
    elif screen_id == "SCR-15":
        nav_flows = [
            "Click \"Save Selected\" → Adds checked items to course deck and returns to Course Structure."
        ]
        display_conds = [
            "Requires role: SME.",
            "Renders AI generated terms, checkbox inputs allow bulk select."
        ]
    elif screen_id == "SCR-16":
        nav_flows = [
            "Click \"Practice Now\" → Redirect to Flashcard Practice (SCR-17)."
        ]
        display_conds = [
            "Requires role: Student.",
            "Displays indicators for cards mastered, learning, and remaining to study."
        ]
    elif screen_id == "SCR-17":
        nav_flows = [
            "Click \"Back to Library\" → Redirect to Flashcard Library (SCR-16)."
        ]
        display_conds = [
            "Requires role: Student.",
            "Cards flip 3D animation when clicking anywhere on the viewport.",
            "Recording button feedback updates internal spaced repetition database schema."
        ]
    elif screen_id == "SCR-18":
        nav_flows = [
            "Click \"Add Question\" → Open Question Detail (SCR-19).",
            "Click \"Import from File\" → Open Question Import (SCR-21).",
            "Click \"AI Generate\" → Open AI Question Staging (SCR-20)."
        ]
        display_conds = [
            "Requires role: SME.",
            "Filter controls support sorting lists by difficulty status and topic tags."
        ]
    elif screen_id == "SCR-19":
        nav_flows = [
            "Click \"Save\" → Saves details and returns to Question Bank (SCR-18)."
        ]
        display_conds = [
            "Requires role: SME.",
            "Requires designating at least one correct choice.",
            "Exposes explanation details to students during subsequent reviews."
        ]
    elif screen_id == "SCR-20":
        nav_flows = [
            "Click \"Approve & Save\" → Save checked items to Question Bank (SCR-18)."
        ]
        display_conds = [
            "Requires role: SME.",
            "SME holds full editing control over choices, questions, and indicators before committing to database."
        ]
    elif screen_id == "SCR-21":
        nav_flows = [
            "Click \"Confirm Import\" → Imports valid lines and redirects to Question Bank (SCR-18)."
        ]
        display_conds = [
            "Requires role: SME.",
            "Accepts standard template structure parameters only, flagging mismatches in real time."
        ]
    elif screen_id == "SCR-22":
        nav_flows = [
            "Click \"Take New Quiz\" → Redirect to New Quiz (SCR-23).",
            "Click \"Review Answers\" → Redirect to Quiz Review (SCR-26)."
        ]
        display_conds = [
            "Requires role: Student.",
            "Displays lists detailing completed items, grades, and pass/fail badges."
        ]
    elif screen_id == "SCR-23":
        nav_flows = [
            "Click \"Start Quiz\" → Redirect to Quiz Taking (SCR-24)."
        ]
        display_conds = [
            "Requires role: Student.",
            "Allows students to dynamically filter and customize testing constraints."
        ]
    elif screen_id == "SCR-24":
        nav_flows = [
            "Click \"Submit Quiz\" → Submits responses and redirects to Quiz Results (SCR-25)."
        ]
        display_conds = [
            "Requires role: Student.",
            "Timer expiration locks the UI and triggers automatic submission.",
            "Unsaved selections cache locally to prevent data loss on network dropouts."
        ]
    elif screen_id == "SCR-25":
        nav_flows = [
            "Click \"Review Answers\" → Redirect to Quiz Review (SCR-26).",
            "Click \"Back to History\" → Redirect to Quiz History (SCR-22)."
        ]
        display_conds = [
            "Requires role: Student.",
            "Displays summary statistics detailing counts correct, incorrect, and skipped."
        ]
    elif screen_id == "SCR-26":
        nav_flows = [
            "Click \"Exit\" → Redirect to Quiz History (SCR-22)."
        ]
        display_conds = [
            "Requires role: Student.",
            "Color codes correct submissions as green, mistakes as red.",
            "Renders SME-designed explanation fields under each item."
        ]
        
    # Set default values if empty
    if not nav_flows:
        nav_flows = [
            "Click main action buttons to save or submit.",
            "Click Back button to return to previous page."
        ]
    if not display_conds:
        role_map = "Student" if "student" in screen_name.lower() or "student" in screen_id.lower() else "SME"
        display_conds = [
            f"Requires user authentication with role: {role_map}.",
            "UI updates instantly upon button interactions."
        ]
        
    return nav_flows, display_conds

def generate_screen_inventory(base_dir):
    print("Analyzing directory:", base_dir)
    inventory_items = []
    
    # Process each configured screen
    for config in SCREENS_CONFIG:
        screen_id = config["id"]
        screen_name = config["name"]
        group = config["group"]
        role = config["role"]
        purpose = config["purpose"]
        
        all_components = []
        parsed_titles = []
        
        # Read from directories
        for s_dir in config["dirs"]:
            full_dir = os.path.join(base_dir, s_dir)
            html_path = os.path.join(full_dir, "code.html")
            
            if os.path.exists(html_path):
                print(f"Parsing HTML for {screen_id}: {html_path}")
                try:
                    with open(html_path, "r", encoding="utf-8") as f:
                        html_content = f.read()
                    
                    if USE_BS4:
                        title, comps = parse_html_bs4(html_content)
                    else:
                        title, comps = parse_html_regex(html_content)
                        
                    parsed_titles.append(title)
                    all_components.extend(comps)
                except Exception as e:
                    print(f"Error parsing {html_path}: {e}")
            else:
                print(f"Warning: HTML file not found: {html_path}")
                
        # Remove duplicate components by name and type
        seen = set()
        unique_components = []
        for c in all_components:
            key = (c["name"].lower(), c["type"].lower())
            if key not in seen:
                seen.add(key)
                unique_components.append(c)
                
        # If no components extracted (e.g. no html file existed yet), add generic ones
        if not unique_components:
            unique_components = [
                {"name": "Page Title", "type": "H1/H2 Title", "desc": f"Displays the header title of the {screen_name} page"},
                {"name": "Back button", "type": "Button", "desc": "Click to return to the previous screen"},
                {"name": "Main Content Area", "type": "Container", "desc": "Main content layout container"}
            ]
            
        nav_flows, display_conds = analyze_navigation_and_conditions(screen_id, screen_name, unique_components)
        
        inventory_items.append({
            "id": screen_id,
            "name": screen_name,
            "group": group,
            "role": role,
            "purpose": purpose,
            "components": unique_components,
            "navigation": nav_flows,
            "conditions": display_conds
        })
        
    return inventory_items

def write_markdown_file(inventory, output_path):
    print("Writing Screen_Inventory.md to", output_path)
    
    # Organize screens by groups
    groups = {}
    for item in inventory:
        group = item["group"]
        if group not in groups:
            groups[group] = []
        groups[group].append(item)
        
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("# EduNexus - Screen Inventory Specification (FDS Section 1)\n\n")
        f.write("> **Source Documents**: Excel Screen List & HTML Interface Mockups.\n")
        f.write("> **Version**: v1.0 (July 2026)\n\n")
        
        f.write("## I. Feature Group Index\n\n")
        for g_idx, (group_name, screens) in enumerate(groups.items(), 1):
            f.write(f"- **{g_idx}. {group_name} Features**\n")
            for scr in screens:
                f.write(f"  - [{scr['id']} - {scr['name']}](#{scr['id'].lower()}---{scr['name'].lower().replace(' ', '-').replace('/', '')})\n")
        f.write("\n---\n\n")
        
        f.write("## II. Detailed Screen Specifications\n\n")
        
        for g_idx, (group_name, screens) in enumerate(groups.items(), 1):
            f.write(f"### {g_idx}. {group_name} Features\n\n")
            
            for scr in screens:
                f.write(f"#### {scr['id']} - {scr['name']}\n\n")
                f.write(f"- **Purpose:** {scr['purpose']}\n")
                f.write(f"- **Role:** `{scr['role']}`\n\n")
                
                f.write("##### a. UI Components\n\n")
                f.write("| Component | Type | Description |\n")
                f.write("| --- | --- | --- |\n")
                for c in scr["components"][:15]: # Limit to first 15 for document formatting
                    f.write(f"| {c['name']} | {c['type']} | {c['desc']} |\n")
                f.write("\n")
                
                f.write("##### b. Navigation Flow\n\n")
                for nav in scr["navigation"]:
                    f.write(f"- {nav}\n")
                f.write("\n")
                
                f.write("##### c. Display Conditions\n\n")
                for cond in scr["conditions"]:
                    f.write(f"- {cond}\n")
                f.write("\n")
                f.write("---\n\n")

    print("Successfully wrote Markdown Screen Inventory.")

def update_docx_file(inventory, docx_template_path, output_docx_path):
    print(f"Reading template {docx_template_path}...")
    if not os.path.exists(docx_template_path):
        print(f"Error: Template file does not exist: {docx_template_path}")
        return False
        
    doc = docx.Document(docx_template_path)
    
    # 1. Update the document title and Group name
    for para in doc.paragraphs:
        if "{{gROUP}}" in para.text:
            para.text = para.text.replace("{{gROUP}}", "EduNexus Learning Platform")
            
    # 2. Modify Section I (Screen Inventory)
    # We find where Section I starts and Section II starts, and replace paragraphs and tables in between.
    elements = doc.element.body
    
    start_idx = -1
    end_idx = -1
    
    # Search elements
    for i, child in enumerate(elements):
        if child.tag.endswith('p'):
            p_text = child.text if hasattr(child, 'text') else ""
            if not p_text:
                # Find paragraphs using docx
                for para in doc.paragraphs:
                    if para._element == child:
                        p_text = para.text
                        break
            
            if "I. Screen Inventory" in p_text:
                start_idx = i
            elif "II. External API Inventory" in p_text:
                end_idx = i
                break
                
    print(f"Start index of Screen Inventory: {start_idx}, End index: {end_idx}")
    
    if start_idx != -1 and end_idx != -1:
        # We delete all elements between start_idx + 1 and end_idx
        # Because indexes shift when deleting, we repeatedly delete the element at start_idx + 1
        num_to_delete = end_idx - start_idx - 1
        print(f"Deleting {num_to_delete} template elements...")
        for _ in range(num_to_delete):
            elements.remove(elements[start_idx + 1])
            
        # Re-initialize the docx wrapper since we manipulated the XML directly
        doc = docx.Document(docx_template_path)
        # Update paragraph text replacements again
        for para in doc.paragraphs:
            if "{{gROUP}}" in para.text:
                para.text = para.text.replace("{{gROUP}}", "EduNexus Learning Platform")
                
        # Now find the paragraph containing "I. Screen Inventory" and insert after it
        start_para = None
        for para in doc.paragraphs:
            if "I. Screen Inventory" in para.text:
                start_para = para
                break
                
        if start_para:
            current_para = start_para
            
            # Organize screens by groups
            groups = {}
            for item in inventory:
                group = item["group"]
                if group not in groups:
                    groups[group] = []
                groups[group].append(item)
                
            for g_idx, (group_name, screens) in enumerate(groups.items(), 1):
                # Add group header paragraph
                current_para = insert_paragraph_after(current_para, f"{g_idx}. {group_name} Features", style='Heading 2')
                
                for scr in screens:
                    # Add screen sub-header
                    current_para = insert_paragraph_after(current_para, f"{scr['id']} - {scr['name']}", style='Heading 3')
                    
                    # Purpose and Role
                    current_para = insert_paragraph_after(current_para, f"Purpose: {scr['purpose']}")
                    current_para = insert_paragraph_after(current_para, f"Role: {scr['role']}")
                    
                    # UI Components sub-heading
                    current_para = insert_paragraph_after(current_para, "a. UI Components", style='Heading 4')
                    
                    # Create UI Components Table
                    current_para = insert_table_after(current_para, doc, scr["components"])
                    
                    # Navigation Flow sub-heading
                    current_para = insert_paragraph_after(current_para, "b. Navigation Flow", style='Heading 4')
                    for nav in scr["navigation"]:
                        current_para = insert_paragraph_after(current_para, f"• {nav}", style='List Bullet')
                        
                    # Display Conditions sub-heading
                    current_para = insert_paragraph_after(current_para, "c. Display Conditions", style='Heading 4')
                    for cond in scr["conditions"]:
                        current_para = insert_paragraph_after(current_para, f"• {cond}", style='List Bullet')
                        
                    # Add space paragraph
                    current_para = insert_paragraph_after(current_para, "")
    else:
        print("Fallback: Could not find exact XML markers. Appending data to Section I.")
        # Fallback approach: just append to doc if markers aren't matched
        doc.add_heading("EduNexus Screen Inventory", level=1)
        for item in inventory:
            doc.add_heading(f"{item['id']} - {item['name']}", level=2)
            doc.add_paragraph(f"Purpose: {item['purpose']}")
            doc.add_paragraph(f"Role: {item['role']}")
            doc.add_heading("UI Components", level=3)
            # Add table
            table = doc.add_table(rows=1, cols=3)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Component'
            hdr_cells[1].text = 'Type'
            hdr_cells[2].text = 'Description'
            for comp in item["components"][:15]:
                row_cells = table.add_row().cells
                row_cells[0].text = comp["name"]
                row_cells[1].text = comp["type"]
                row_cells[2].text = comp["desc"]
            doc.add_heading("Navigation Flow", level=3)
            for nav in item["navigation"]:
                doc.add_paragraph(nav, style='List Bullet')
            doc.add_heading("Display Conditions", level=3)
            for cond in item["conditions"]:
                doc.add_paragraph(cond, style='List Bullet')
                
    doc.save(output_docx_path)
    print(f"Successfully saved updated DOCX file to {output_docx_path}")
    return True

def insert_paragraph_after(paragraph, text, style=None):
    new_p = paragraph._element.getnext()
    if new_p is None:
        # Paragraph is the last element
        p = paragraph.part.add_paragraph(text, style=style)
    else:
        p = paragraph.part.add_paragraph(text, style=style)
        # Move it to correct position XML-wise
        paragraph._element.addnext(p._element)
    return p

def insert_table_after(paragraph, doc, components):
    table = doc.add_table(rows=1, cols=3)
    # Apply standard table formatting if necessary
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Component'
    hdr_cells[1].text = 'Type'
    hdr_cells[2].text = 'Description'
    
    for comp in components[:15]:
        row_cells = table.add_row().cells
        row_cells[0].text = comp["name"]
        row_cells[1].text = comp["type"]
        row_cells[2].text = comp["desc"]
        
    # Insert table after paragraph XML-wise
    paragraph._element.addnext(table._element)
    
    # We return a dummy paragraph after the table to continue insertion
    p = paragraph.part.add_paragraph("")
    table._element.addnext(p._element)
    return p

if __name__ == "__main__":
    current_dir = os.path.dirname(os.path.abspath(__file__))
    print("EduNexus FDS Generator Tool")
    print("---------------------------")
    print("Current Repo Dir:", current_dir)
    
    # Generate Inventory
    inventory = generate_screen_inventory(current_dir)
    
    # Write Markdown
    md_output_path = os.path.join(current_dir, "Screen_Inventory.md")
    write_markdown_file(inventory, md_output_path)
    
    # Update DOCX template
    template_docx_path = r"d:\Gx_3-Functional-Design-Spec.docx"
    output_docx_path = os.path.join(current_dir, "Gx_3-Functional-Design-Spec-EduNexus.docx")
    
    success = False
    if os.path.exists(template_docx_path):
        success = update_docx_file(inventory, template_docx_path, output_docx_path)
    else:
        # Fallback to look in current dir or parents
        local_template = os.path.join(current_dir, "Gx_3-Functional-Design-Spec.docx")
        if os.path.exists(local_template):
            success = update_docx_file(inventory, local_template, output_docx_path)
        else:
            print(f"Warning: Template docx not found at {template_docx_path} or {local_template}. Skipping Word document generation.")
            
    print("---------------------------")
    if success:
        print("Successfully generated Screen_Inventory.md and Gx_3-Functional-Design-Spec-EduNexus.docx!")
    else:
        print("Successfully generated Screen_Inventory.md only. Please run this script on a machine with python-docx and access to d:\\Gx_3-Functional-Design-Spec.docx to get the Word document.")
